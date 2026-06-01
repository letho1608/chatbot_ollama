"""Generate offline PNG diagrams and the Vietnamese Tysor user guide DOCX."""

from __future__ import annotations

import re
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
DIAGRAM_DIR = DOCS_DIR / "diagrams"
GUIDE_PATH = DOCS_DIR / "HUONG_DAN_SU_DUNG.md"
DOCX_PATH = DOCS_DIR / "Tysor_Huong_Dan_Su_Dung.docx"

CANVAS_WIDTH = 1800
MARGIN_X = 80
TOP_Y = 145
LAYER_GAP = 95
BOX_GAP = 34
BOX_MIN_HEIGHT = 112

COLORS = {
    "background": "#f8fafc",
    "text": "#172033",
    "muted": "#516071",
    "line": "#64748b",
    "browser": "#dbeafe",
    "server": "#d1fae5",
    "agent": "#ccfbf1",
    "data": "#fef3c7",
    "external": "#ede9fe",
    "security": "#fee2e2",
    "admin": "#f3e8ff",
    "default": "#e2e8f0",
}


@dataclass(frozen=True)
class Node:
    key: str
    label: str
    layer: int
    group: str = "default"


@dataclass(frozen=True)
class Edge:
    source: str
    target: str
    label: str = ""


@dataclass(frozen=True)
class Diagram:
    filename: str
    title: str
    subtitle: str
    nodes: tuple[Node, ...]
    edges: tuple[Edge, ...]
    note: str


DIAGRAMS = (
    Diagram(
        "01-system-architecture.png",
        "1. Tysor System Architecture",
        "FastAPI web application with local Ollama runtime and optional public tunnel",
        (
            Node("browser", "Web Browser\nDashboard | Login | Chat | Admin", 0, "browser"),
            Node("tunnel", "Cloudflare Tunnel\noptional public URL", 0, "external"),
            Node("fastapi", "FastAPI Server\nPages | REST API | SSE streaming", 1, "server"),
            Node("sqlite", "SQLite\ndata/chatbot.db", 2, "data"),
            Node("agent", "Agent Framework\nSafety | Memory | Validator | Monitor | Queue", 2, "agent"),
            Node("hybrid", "Hybrid Retrieval\nSession RAG | GraphRAG | Web fallback", 3, "agent"),
            Node("ollama", "Ollama 127.0.0.1:11434\nchat | embeddings | pull | tags", 4, "server"),
            Node("dataset", "AISecKG Dataset\nCSV triples + entity info", 4, "data"),
            Node("web", "Trusted Cyber Security Websites\nallowed-domain fallback", 4, "external"),
        ),
        (
            Edge("tunnel", "fastapi", "public HTTP"),
            Edge("browser", "fastapi", "HTTP / SSE"),
            Edge("fastapi", "sqlite", "SQLAlchemy"),
            Edge("fastapi", "agent", "orchestrate"),
            Edge("agent", "hybrid", "retrieve context"),
            Edge("agent", "ollama", "/api/chat"),
            Edge("hybrid", "ollama", "/api/embeddings"),
            Edge("hybrid", "dataset", "GraphRAG"),
            Edge("hybrid", "web", "fallback"),
        ),
        "The tunnel is optional. Local use only needs the browser, FastAPI and Ollama.",
    ),
    Diagram(
        "02-authentication-flow.png",
        "2. Authentication Flow",
        "JWT cookie and localStorage token with role-based redirect",
        (
            Node("form", "Browser Form\n/register or /login", 0, "browser"),
            Node("auth", "FastAPI Auth API\nPOST /api/auth/register\nPOST /api/auth/login", 1, "server"),
            Node("user", "SQLite users\nunique username + email", 2, "data"),
            Node("crypto", "Password Handling\nhash on register\nverify on login", 3, "security"),
            Node("audit", "SQLite audit_logs\nregister | login | login_failed", 3, "data"),
            Node("jwt", "JWT Response\nHttpOnly cookie + JSON token", 4, "security"),
            Node("storage", "Browser\nlocalStorage ollama_token", 5, "browser"),
            Node("me", "GET /api/auth/me\nvalidate active user", 6, "server"),
            Node("route", "Role Redirect\nuser -> /chat\nadmin -> /admin", 7, "browser"),
        ),
        (
            Edge("form", "auth"),
            Edge("auth", "user", "lookup / insert"),
            Edge("user", "crypto"),
            Edge("crypto", "audit", "write event"),
            Edge("crypto", "jwt", "create token"),
            Edge("jwt", "storage"),
            Edge("storage", "me", "Bearer token"),
            Edge("me", "route"),
        ),
        "On first startup, the server creates admin/admin123 only if no admin exists.",
    ),
    Diagram(
        "03-chat-streaming-flow.png",
        "3. Chat Streaming Flow",
        "Validated input, queued Ollama request, redacted SSE output and audit trail",
        (
            Node("browser", "Browser\nPOST /api/chat/stream", 0, "browser"),
            Node("input", "Agent process_input()\nsafety | injection | topic | redaction", 1, "security"),
            Node("save_user", "SQLite\nsave original user message\nauto-name conversation", 2, "data"),
            Node("build", "Build Ollama Messages\nbase prompt | memory | hybrid RAG | history", 3, "agent"),
            Node("limits", "Server Limits\nadaptive options | rate limit 20/min", 4, "server"),
            Node("queue", "Auto-scaling Queue\n2..6 workers | Ollama semaphore=1", 5, "agent"),
            Node("ollama", "Ollama /api/chat\nstream response chunks", 6, "server"),
            Node("redact", "Streaming Redaction\nsanitize secrets + PII before SSE", 7, "security"),
            Node("render", "Browser\nrender Markdown progressively", 8, "browser"),
            Node("save_asst", "SQLite + Agent Post-process\nsave response | validate | memory | audit", 9, "data"),
        ),
        (
            Edge("browser", "input"),
            Edge("input", "save_user"),
            Edge("save_user", "build"),
            Edge("build", "limits"),
            Edge("limits", "queue"),
            Edge("queue", "ollama"),
            Edge("ollama", "redact"),
            Edge("redact", "render", "SSE chunks"),
            Edge("redact", "save_asst", "final redacted text"),
        ),
        "The browser receives redacted chunks. The saved assistant response is also redacted.",
    ),
    Diagram(
        "04-hybrid-rag-flow.png",
        "4. Hybrid RAG Flow",
        "Session upload plus GraphRAG and trusted-domain web fallback",
        (
            Node("upload", "POST /api/rag/upload-file", 0, "browser"),
            Node("extension", "Validate Extension\ntext | config | script | docx | pdf", 1, "security"),
            Node("extract", "Extract Text\nminimum 50 characters", 2, "server"),
            Node("keywords", "Security Content Check\nminimum 2 Cyber Security keywords", 3, "security"),
            Node("chunk", "Chunk Text\n500 words | overlap 80", 4, "agent"),
            Node("embed", "Ollama /api/embeddings\nnomic-embed-text", 5, "server"),
            Node("session", "Per-user session_store\nin-memory documents + vectors", 6, "data"),
            Node("query", "Chat Query", 7, "browser"),
            Node("session_search", "Session Cosine Similarity", 8, "agent"),
            Node("graph", "GraphRAG\nAISecKG dataset", 8, "data"),
            Node("merge", "Merge Retrieval Results", 9, "agent"),
            Node("confidence", "Confident Result?\nFresh-web query?", 10, "security"),
            Node("web", "Trusted-domain Web Fallback", 11, "external"),
            Node("prompt", "Inject Context Into System Prompt", 12, "agent"),
        ),
        (
            Edge("upload", "extension"),
            Edge("extension", "extract", "allowed"),
            Edge("extract", "keywords"),
            Edge("keywords", "chunk", "accepted"),
            Edge("chunk", "embed"),
            Edge("embed", "session"),
            Edge("session", "session_search"),
            Edge("query", "session_search"),
            Edge("query", "graph"),
            Edge("session_search", "merge"),
            Edge("graph", "merge"),
            Edge("merge", "confidence"),
            Edge("confidence", "prompt", "yes"),
            Edge("confidence", "web", "fallback"),
            Edge("web", "prompt"),
        ),
        "Files uploaded from the chat UI are session-only and disappear after server restart.",
    ),
    Diagram(
        "05-safety-redaction-flow.png",
        "5. Safety And Redaction Flow",
        "Input checks before Ollama and output redaction before browser delivery",
        (
            Node("query", "User Query", 0, "browser"),
            Node("content", "1. Content Safety\nblocked categories", 1, "security"),
            Node("injection", "2. Prompt Injection Detection\njailbreak patterns", 2, "security"),
            Node("topic", "3. Security Relevance\nskip for greeting/social", 3, "security"),
            Node("sensitive", "4. sanitize_sensitive()\nsecrets + PII placeholders", 4, "security"),
            Node("ollama", "Ollama /api/chat", 5, "server"),
            Node("stream", "Stream Chunks", 6, "server"),
            Node("redact", "sanitize_sensitive()\nbefore SSE delivery", 7, "security"),
            Node("browser", "Browser Render", 8, "browser"),
            Node("post", "Agent process_response()", 8, "agent"),
            Node("validate", "Validator + Output Safety", 9, "security"),
            Node("memory", "Memory Extraction + Audit Log", 10, "data"),
        ),
        (
            Edge("query", "content"),
            Edge("content", "injection", "pass"),
            Edge("injection", "topic", "pass"),
            Edge("topic", "sensitive", "pass"),
            Edge("sensitive", "ollama"),
            Edge("ollama", "stream"),
            Edge("stream", "redact"),
            Edge("redact", "browser", "SSE"),
            Edge("stream", "post", "final response"),
            Edge("post", "validate"),
            Edge("validate", "memory"),
        ),
        "Blocked input returns HTTP 400. Sensitive values are replaced with placeholders.",
    ),
    Diagram(
        "06-admin-flow.png",
        "6. Admin Panel Flow",
        "Role-protected operations exposed by the admin dashboard",
        (
            Node("admin", "Admin Browser\n/admin", 0, "browser"),
            Node("auth", "GET /api/auth/me\nrequire role=admin", 1, "security"),
            Node("dashboard", "Dashboard\nstats + activity", 2, "admin"),
            Node("users", "Users\ndetail | role | toggle | delete", 2, "admin"),
            Node("models", "Models\ncatalog | pull | delete", 2, "admin"),
            Node("convs", "Conversations\nbrowse all users", 2, "admin"),
            Node("audit", "Audit Logs\nfilter events", 3, "data"),
            Node("system", "System Health\nOllama | RAM | disk | process | AI log", 3, "server"),
            Node("tunnel", "Cloudflare Tunnel\nstatus | start | stop", 4, "external"),
        ),
        (
            Edge("admin", "auth"),
            Edge("auth", "dashboard"),
            Edge("auth", "users"),
            Edge("auth", "models"),
            Edge("auth", "convs"),
            Edge("dashboard", "audit"),
            Edge("models", "system"),
            Edge("convs", "audit"),
            Edge("system", "tunnel"),
        ),
        "Do not expose the public tunnel while the default admin password is still active.",
    ),
    Diagram(
        "07-agent-framework.png",
        "7. Agent Framework Architecture",
        "Orchestration stages used for each chat request",
        (
            Node("input", "process_input()\nvalidate + redact", 0, "security"),
            Node("build", "build_ollama_messages()", 1, "agent"),
            Node("prompt", "Base + Custom\nSystem Prompt", 2, "server"),
            Node("memory", "Persistent\nUser Memories", 2, "data"),
            Node("rag", "Hybrid RAG\nContext", 2, "agent"),
            Node("history", "Trimmed\nConversation History", 2, "data"),
            Node("messages", "Ollama Messages", 3, "agent"),
            Node("adaptive", "Adaptive Options\n512 | 2048 | 4096 tokens", 4, "server"),
            Node("response", "process_response()", 5, "agent"),
            Node("validator", "Response Validator\nloop | consistency | confidence", 6, "security"),
            Node("extract", "Memory Extraction", 6, "data"),
            Node("monitor", "Audit + AI Log", 6, "data"),
        ),
        (
            Edge("input", "build"),
            Edge("build", "prompt"),
            Edge("build", "memory"),
            Edge("build", "rag"),
            Edge("build", "history"),
            Edge("prompt", "messages"),
            Edge("memory", "messages"),
            Edge("rag", "messages"),
            Edge("history", "messages"),
            Edge("messages", "adaptive"),
            Edge("adaptive", "response"),
            Edge("response", "validator"),
            Edge("response", "extract"),
            Edge("response", "monitor"),
        ),
        "Hybrid RAG may combine uploaded files, AISecKG graph facts and trusted web context.",
    ),
    Diagram(
        "08-database-schema.png",
        "8. SQLite Database Schema",
        "Persistent SQLAlchemy tables in data/chatbot.db",
        (
            Node("users", "users\nid PK | username UQ | email UQ\npassword_hash | role | is_active | created_at", 0, "data"),
            Node("preferences", "user_preferences\nid PK | user_id FK UQ\nlanguage | theme | max_tokens | temperature | extra", 1, "data"),
            Node("conversations", "conversations\nid string PK | user_id FK\ntitle | created_at | updated_at", 1, "data"),
            Node("memories", "memories\nid PK | user_id FK | key\nvalue | source | confidence | timestamps", 1, "data"),
            Node("audit", "audit_logs\nid PK | user_id FK nullable\nevent | detail JSON | ip_address | created_at", 1, "data"),
            Node("messages", "messages\nid PK | conversation_id FK\nrole | content | timestamp", 2, "data"),
            Node("documents", "rag_documents\nid PK | user_id FK\nfilename | content | chunk_count | created_at", 2, "data"),
            Node("chunks", "rag_chunks\nid PK | document_id FK\nchunk_index | content | embedding text", 3, "data"),
        ),
        (
            Edge("users", "preferences", "0..1"),
            Edge("users", "conversations", "1..N"),
            Edge("users", "memories", "1..N"),
            Edge("users", "audit", "1..N"),
            Edge("users", "documents", "1..N"),
            Edge("conversations", "messages", "1..N"),
            Edge("documents", "chunks", "1..N"),
        ),
        "The current chat upload API uses in-memory session_store; it does not persist UI uploads to rag_documents.",
    ),
)


def _font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    windows_fonts = Path("C:/Windows/Fonts")
    candidates = (
        ("consolab.ttf", "consola.ttf")
        if mono
        else ("segoeuib.ttf", "arialbd.ttf") if bold else ("segoeui.ttf", "arial.ttf")
    )
    for candidate in candidates:
        path = windows_fonts / candidate
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    left, _, right, _ = draw.textbbox((0, 0), text, font=font)
    return right - left


def _wrap_line(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and _text_width(draw, candidate, font) > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [text]


def _wrapped_lines(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont,
    max_width: int,
) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        lines.extend(_wrap_line(draw, raw, font, max_width))
    return lines


def _arrow_head(draw: ImageDraw.ImageDraw, x: int, y: int, direction: str) -> None:
    size = 12
    if direction == "down":
        points = ((x, y), (x - size, y - size * 2), (x + size, y - size * 2))
    elif direction == "up":
        points = ((x, y), (x - size, y + size * 2), (x + size, y + size * 2))
    elif direction == "right":
        points = ((x, y), (x - size * 2, y - size), (x - size * 2, y + size))
    else:
        points = ((x, y), (x + size * 2, y - size), (x + size * 2, y + size))
    draw.polygon(points, fill=COLORS["line"])


def _draw_edge(
    draw: ImageDraw.ImageDraw,
    source: tuple[int, int, int, int],
    target: tuple[int, int, int, int],
    label: str,
    edge_font: ImageFont.ImageFont,
) -> None:
    sx1, sy1, sx2, sy2 = source
    tx1, ty1, tx2, ty2 = target
    source_cx, source_cy = (sx1 + sx2) // 2, (sy1 + sy2) // 2
    target_cx, target_cy = (tx1 + tx2) // 2, (ty1 + ty2) // 2

    if sy2 < ty1:
        start = (source_cx, sy2)
        end = (target_cx, ty1)
        mid_y = (start[1] + end[1]) // 2
        points = (start, (start[0], mid_y), (end[0], mid_y), end)
        direction = "down"
    elif ty2 < sy1:
        start = (source_cx, sy1)
        end = (target_cx, ty2)
        mid_y = (start[1] + end[1]) // 2
        points = (start, (start[0], mid_y), (end[0], mid_y), end)
        direction = "up"
    elif sx2 < tx1:
        start = (sx2, source_cy)
        end = (tx1, target_cy)
        mid_x = (start[0] + end[0]) // 2
        points = (start, (mid_x, start[1]), (mid_x, end[1]), end)
        direction = "right"
    else:
        start = (sx1, source_cy)
        end = (tx2, target_cy)
        mid_x = (start[0] + end[0]) // 2
        points = (start, (mid_x, start[1]), (mid_x, end[1]), end)
        direction = "left"

    draw.line(points, fill=COLORS["line"], width=4, joint="curve")
    _arrow_head(draw, end[0], end[1], direction)

    if label:
        point_a = points[len(points) // 2 - 1]
        point_b = points[len(points) // 2]
        lx = (point_a[0] + point_b[0]) // 2
        ly = (point_a[1] + point_b[1]) // 2
        bbox = draw.textbbox((lx, ly), label, font=edge_font, anchor="mm")
        pad = 5
        draw.rounded_rectangle(
            (bbox[0] - pad, bbox[1] - pad, bbox[2] + pad, bbox[3] + pad),
            radius=5,
            fill="#ffffff",
            outline="#cbd5e1",
        )
        draw.text((lx, ly), label, fill=COLORS["muted"], font=edge_font, anchor="mm")


def render_diagram(diagram: Diagram, destination: Path) -> None:
    draft = Image.new("RGB", (CANVAS_WIDTH, 3000), COLORS["background"])
    draw = ImageDraw.Draw(draft)
    title_font = _font(46, bold=True)
    subtitle_font = _font(24)
    node_font = _font(24)
    edge_font = _font(18)
    note_font = _font(22)

    by_layer: dict[int, list[Node]] = {}
    for node in diagram.nodes:
        by_layer.setdefault(node.layer, []).append(node)

    boxes: dict[str, tuple[int, int, int, int]] = {}
    y = TOP_Y
    for layer in sorted(by_layer):
        nodes = by_layer[layer]
        count = len(nodes)
        available_width = CANVAS_WIDTH - 2 * MARGIN_X - BOX_GAP * max(count - 1, 0)
        box_width = min(420, available_width // count)
        total_width = box_width * count + BOX_GAP * max(count - 1, 0)
        x = (CANVAS_WIDTH - total_width) // 2

        prepared: list[tuple[Node, list[str], int]] = []
        row_height = BOX_MIN_HEIGHT
        for node in nodes:
            lines = _wrapped_lines(draw, node.label, node_font, box_width - 34)
            height = max(BOX_MIN_HEIGHT, len(lines) * 32 + 34)
            prepared.append((node, lines, height))
            row_height = max(row_height, height)

        for node, _, _ in prepared:
            boxes[node.key] = (x, y, x + box_width, y + row_height)
            x += box_width + BOX_GAP
        y += row_height + LAYER_GAP

    final_height = y + 105
    image = draft.crop((0, 0, CANVAS_WIDTH, final_height))
    draw = ImageDraw.Draw(image)

    draw.text((MARGIN_X, 42), diagram.title, fill=COLORS["text"], font=title_font)
    draw.text((MARGIN_X, 100), diagram.subtitle, fill=COLORS["muted"], font=subtitle_font)

    for edge in diagram.edges:
        _draw_edge(draw, boxes[edge.source], boxes[edge.target], edge.label, edge_font)

    for node in diagram.nodes:
        box = boxes[node.key]
        draw.rounded_rectangle(
            box,
            radius=18,
            fill=COLORS.get(node.group, COLORS["default"]),
            outline="#475569",
            width=3,
        )
        lines = _wrapped_lines(draw, node.label, node_font, box[2] - box[0] - 34)
        line_height = 32
        block_height = len(lines) * line_height
        text_y = (box[1] + box[3] - block_height) // 2 + line_height // 2
        for index, line in enumerate(lines):
            draw.text(
                ((box[0] + box[2]) // 2, text_y + index * line_height),
                line,
                fill=COLORS["text"],
                font=node_font,
                anchor="mm",
            )

    note_y = final_height - 72
    draw.rounded_rectangle(
        (MARGIN_X, note_y - 18, CANVAS_WIDTH - MARGIN_X, final_height - 24),
        radius=10,
        fill="#ffffff",
        outline="#cbd5e1",
    )
    wrapped_note = textwrap.wrap(f"Note: {diagram.note}", width=135)
    for index, line in enumerate(wrapped_note[:2]):
        draw.text((MARGIN_X + 18, note_y + index * 26), line, fill=COLORS["muted"], font=note_font)

    destination.parent.mkdir(parents=True, exist_ok=True)
    image.save(destination, "PNG", optimize=True)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _add_toc(document: Document) -> None:
    document.add_heading("Mục lục", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Mở Word và chọn Update Field để cập nhật mục lục."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$")


def _add_inline(paragraph, text: str) -> None:
    for token in filter(None, INLINE_RE.split(text)):
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(16, 103, 122)
        elif token.startswith("[") and "](" in token:
            label, url = token[1:-1].split("](", 1)
            paragraph.add_run(f"{label} ({url})")
        else:
            paragraph.add_run(token)


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        if not TABLE_SEPARATOR_RE.match(line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        index += 1
    return rows, index


def _add_markdown(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_first_title = False

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph(style="Code Block")
                paragraph.add_run("\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not line.strip() or line.strip() == "---":
            index += 1
            continue

        if line.startswith("# "):
            if not skipped_first_title:
                skipped_first_title = True
            else:
                document.add_heading(line[2:].strip(), level=1)
            index += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            index += 1
            continue

        if line.startswith("#### "):
            document.add_heading(line[5:].strip(), level=3)
            index += 1
            continue

        image_match = IMAGE_RE.match(line.strip())
        if image_match:
            alt, rel_path = image_match.groups()
            image_path = markdown_path.parent / rel_path
            with Image.open(image_path) as image:
                ratio = image.height / image.width
            if ratio * 6.35 > 8.0:
                document.add_picture(str(image_path), height=Inches(8.0))
            else:
                document.add_picture(str(image_path), width=Inches(6.35))
            caption = document.add_paragraph(style="Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(alt)
            index += 1
            continue

        if line.strip().startswith("|"):
            rows, index = _parse_table(lines, index)
            if not rows:
                continue
            columns = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=columns)
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for column_index in range(columns):
                    value = row[column_index] if column_index < len(row) else ""
                    _set_cell_text(table.cell(row_index, column_index), value, bold=row_index == 0)
            document.add_paragraph()
            continue

        if re.match(r"^\s*-\s+", line):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, re.sub(r"^\s*-\s+", "", line))
            index += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            paragraph = document.add_paragraph(style="List Number")
            _add_inline(paragraph, re.sub(r"^\s*\d+\.\s+", "", line))
            index += 1
            continue

        paragraph = document.add_paragraph()
        _add_inline(paragraph, line)
        index += 1


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in (
        ("Title", 24, "0f766e"),
        ("Heading 1", 16, "0f766e"),
        ("Heading 2", 13, "155e75"),
        ("Heading 3", 11, "334155"),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(15, 118, 110)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TYSOR | Hướng dẫn cài đặt và sử dụng")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(100, 116, 139)
    _add_page_number(section.footer.paragraphs[0])


def build_docx() -> None:
    document = Document()
    _configure_document(document)
    document.core_properties.title = "Hướng dẫn cài đặt và sử dụng Tysor"
    document.core_properties.subject = "Tysor Cyber Security AI Chatbot"
    document.core_properties.author = "Tysor project"

    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("HƯỚNG DẪN CÀI ĐẶT\nVÀ SỬ DỤNG TYSOR")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Cyber Security AI Chatbot\n").bold = True
    subtitle.add_run("Tài liệu vận hành chi tiết và diagram hệ thống")

    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Được đối chiếu với mã nguồn hiện tại của dự án.").italic = True

    document.add_page_break()
    _add_toc(document)
    document.add_page_break()
    _add_markdown(document, GUIDE_PATH)

    # Ensure the last section inherits the standard page setup.
    final_section = document.add_section(WD_SECTION.CONTINUOUS)
    final_section.top_margin = Cm(1.8)
    final_section.bottom_margin = Cm(1.6)
    final_section.left_margin = Cm(1.8)
    final_section.right_margin = Cm(1.8)

    document.save(DOCX_PATH)


def main() -> None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    for diagram in DIAGRAMS:
        destination = DIAGRAM_DIR / diagram.filename
        render_diagram(diagram, destination)
        print(f"generated {destination.relative_to(ROOT)}")
    build_docx()
    print(f"generated {DOCX_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
