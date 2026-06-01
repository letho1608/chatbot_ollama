import json
import math
import os
import re
import uuid
from io import BytesIO
from typing import List, Tuple, Optional

import httpx
from sqlalchemy.orm import Session

from core.database import Document, DocumentChunk
from core.config import SECURITY_KEYWORDS

OLLAMA_BASE = "http://127.0.0.1:11434"
EMBED_MODEL = "nomic-embed-text"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 80
RAG_RESULTS = 4
WEB_FETCH_SESSION_MIN_SCORE = float(os.getenv("WEB_FETCH_SESSION_MIN_SCORE", "0.24"))
WEB_FETCH_GRAPH_MIN_SCORE = float(os.getenv("WEB_FETCH_GRAPH_MIN_SCORE", "3.0"))
WEB_FRESHNESS_TERMS = (
    "latest", "recent", "current", "today", "yesterday", "this week",
    "newest", "news", "advisory", "mới nhất", "gần đây", "hôm nay",
    "hiện tại", "vừa", "cảnh báo",
)


# In-memory session-only document store: {user_id: [doc, ...]}
# doc = {"id": str, "filename": str, "chunks": [{"content": str, "embedding": list}, ...]}
session_store: dict[int, list[dict]] = {}


def is_security_content(text: str) -> tuple[bool, str]:
    if len(text.strip()) < 50:
        return False, "Nội dung quá ngắn (tối thiểu 50 ký tự)"
    text_lower = text.lower()
    matches = sum(1 for kw in SECURITY_KEYWORDS if kw in text_lower)
    if matches < 2:
        return False, f"Nội dung không liên quan đến Cyber Security (chỉ tìm thấy {matches}/2 từ khóa)"
    return True, ""


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    text = re.sub(r'\s+', ' ', text).strip()
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(' '.join(words[start:end]))
        if end >= len(words):
            break
        start = end - overlap
    return chunks


def get_embedding(text: str, model: str = EMBED_MODEL) -> List[float]:
    try:
        with httpx.Client(timeout=30) as client:
            resp = client.post(f"{OLLAMA_BASE}/api/embeddings", json={"model": model, "prompt": text})
            if resp.status_code == 200:
                return resp.json().get("embedding", [])
    except (httpx.HTTPError, ValueError):
        return []
    return []


def cosine_similarity(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def store_document_sync(user_id: int, filename: str, content: str, db: Session) -> Document:
    doc = Document(user_id=user_id, filename=filename, content=content)
    db.add(doc)
    db.flush()

    chunks = chunk_text(content)
    doc.chunk_count = len(chunks)

    for i, chunk in enumerate(chunks):
        emb = get_embedding(chunk)
        emb_json = json.dumps(emb) if emb else None
        c = DocumentChunk(document_id=doc.id, chunk_index=i, content=chunk, embedding=emb_json)
        db.add(c)

    db.flush()
    return doc


def retrieve(user_id: int, query: str, db: Session, top_k: int = RAG_RESULTS) -> List[Tuple[str, float]]:
    q_emb = get_embedding(query)
    if not q_emb:
        return []

    chunks = db.query(DocumentChunk).join(Document).filter(
        Document.user_id == user_id,
        DocumentChunk.embedding.isnot(None),
    ).all()

    scored = []
    for c in chunks:
        c_emb = json.loads(c.embedding)
        if c_emb:
            score = cosine_similarity(q_emb, c_emb)
            scored.append((c.content, score, c.document.filename))

    scored.sort(key=lambda x: x[1], reverse=True)
    return [(content, score) for content, score, _ in scored[:top_k]]


def format_rag_context(results: List[Tuple[str, float]]) -> str:
    if not results:
        return ""
    parts = []
    for content, score in results:
        parts.append(content)
    return "Relevant context (uploaded documents, knowledge graph facts, and trusted web sources):\n" + "\n---\n".join(parts)


# ─── Session-only (in-memory) RAG ─────────────────────────────────────────────

def store_session_document(user_id: int, filename: str, content: str) -> dict:
    chunks = chunk_text(content)
    embedded_chunks = []
    for chunk in chunks:
        emb = get_embedding(chunk)
        embedded_chunks.append({"content": chunk, "embedding": emb})
    doc = {
        "id": str(uuid.uuid4()),
        "filename": filename,
        "chunks": embedded_chunks,
        "chunk_count": len(chunks),
    }
    if user_id not in session_store:
        session_store[user_id] = []
    session_store[user_id].append(doc)
    return doc


def get_session_docs(user_id: int) -> list[dict]:
    return session_store.get(user_id, [])


def delete_session_doc(user_id: int, doc_id: str) -> bool:
    docs = session_store.get(user_id, [])
    for i, d in enumerate(docs):
        if d["id"] == doc_id:
            docs.pop(i)
            return True
    return False


def session_retrieve(user_id: int, query: str, top_k: int = RAG_RESULTS) -> List[Tuple[str, float]]:
    q_emb = get_embedding(query)
    if not q_emb or user_id not in session_store:
        return []
    scored = []
    for doc in session_store[user_id]:
        for chunk in doc["chunks"]:
            if chunk["embedding"]:
                score = cosine_similarity(q_emb, chunk["embedding"])
                scored.append((chunk["content"], score))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored[:top_k]


def graphrag_retrieve(query: str, top_k: int = RAG_RESULTS) -> List[Tuple[str, float]]:
    try:
        from core.graphrag import retrieve_graph_context
        return retrieve_graph_context(query, top_k=top_k)
    except Exception:
        return []


def web_retrieve(query: str, top_k: int = RAG_RESULTS) -> List[Tuple[str, float]]:
    try:
        from core.web_fetch import retrieve_web_context
        return retrieve_web_context(query, top_k=top_k)
    except Exception:
        return []


def query_needs_fresh_web(query: str) -> bool:
    query_lower = query.lower()
    return any(term in query_lower for term in WEB_FRESHNESS_TERMS)


def has_confident_rag_hit(results: List[Tuple[str, float]], query: str = "") -> bool:
    needs_fresh_web = query_needs_fresh_web(query)
    for content, score in results:
        if content.startswith("Web context:"):
            continue
        if needs_fresh_web and content.startswith("Graph fact:"):
            continue
        if content.startswith("Graph fact:") and score >= WEB_FETCH_GRAPH_MIN_SCORE:
            return True
        if not content.startswith("Graph fact:") and score >= WEB_FETCH_SESSION_MIN_SCORE:
            return True
    return False


def hybrid_retrieve(user_id: int, query: str, top_k: int = RAG_RESULTS, include_web: bool = True) -> List[Tuple[str, float]]:
    session_results = session_retrieve(user_id, query, top_k=top_k)
    graph_results = graphrag_retrieve(query, top_k=top_k)
    results = session_results + graph_results
    if include_web and not has_confident_rag_hit(results, query):
        results.extend(web_retrieve(query, top_k=top_k))
    return results


SESSION_ALLOWED_EXTENSIONS = {".txt", ".md", ".csv", ".log", ".json", ".yaml", ".yml", ".conf", ".cfg", ".ini", ".xml", ".ps1", ".sh", ".py", ".bat", ".docx", ".pdf"}


def is_allowed_extension(filename: str) -> bool:
    _, ext = os.path.splitext(filename.lower())
    return ext in SESSION_ALLOWED_EXTENSIONS


def extract_text_from_bytes(data: bytes, filename: str) -> Optional[str]:
    ext = os.path.splitext(filename.lower())[1]
    try:
        if ext in (".txt", ".md", ".csv", ".log", ".json", ".yaml", ".yml", ".conf", ".cfg", ".ini", ".xml", ".ps1", ".sh", ".py", ".bat"):
            return data.decode("utf-8", errors="replace").strip()
        elif ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(BytesIO(data))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paras).strip()
        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(data))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    pages.append(text.strip())
            return "\n".join(pages).strip()
    except Exception:
        return None
    return None
